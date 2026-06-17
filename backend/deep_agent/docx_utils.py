import re
import zipfile
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

def extract_comments_from_docx(file_stream: BytesIO):
    """
    Extracts comments and their context from a DOCX file.
    Returns a list of dicts: {id, comment, report_text, context_before, context_after}
    """
    comments = []
    
    try:
        with zipfile.ZipFile(file_stream) as docx_zip:
            # 1. Read comments.xml
            try:
                comments_xml = docx_zip.read('word/comments.xml')
                comments_root = etree.fromstring(comments_xml)
            except KeyError:
                return [] # No comments found

            # Namespace map
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            # Parse comments
            comments_map = {}
            for comment in comments_root.xpath('//w:comment', namespaces=ns):
                comment_id = comment.get(f"{{{ns['w']}}}id")
                # Extract text from all paragraphs within the comment
                text_parts = comment.xpath('.//w:t/text()', namespaces=ns)
                comment_text = "".join(text_parts)
                comments_map[comment_id] = comment_text

            # 2. Read document.xml to find ranges
            document_xml = docx_zip.read('word/document.xml')
            doc_root = etree.fromstring(document_xml)
            
            # Helper to get all text from document
            # We need a way to map XML nodes to linear text or traverse efficiently
            # A simple approach for "context" is collecting all text nodes and finding indices, 
            # but mapping back to XML ranges is tricky.
            
            # Alternative: Traverse the document tree.
            # When we hit commentRangeStart, we mark start.
            # When we hit commentRangeEnd, we mark end.
            
            # Let's linearize the document text but keep track of node references or just text content?
            # Requirement: "highlighted text" (between range) + 100 words before + 100 words after.
            
            # Strategy:
            # 1. Extract all text nodes with their parents, keeping order.
            # 2. Identify where comment starts/ends occur in this inspection order.
            
            full_text_content = ""
            node_positions = [] # (type, id, text_length_so_far)
            
            # Robust traversal using iterwalk or xpath might be complex. 
            # formatting checks: //w:t | //w:commentRangeStart | //w:commentRangeEnd
            
            # Robust traversal using iterwalk or xpath might be complex. 
            # formatting checks: //w:t | //w:commentRangeStart | //w:commentRangeEnd
            
            elements = doc_root.xpath('//w:t | //w:commentRangeStart | //w:commentRangeEnd', namespaces=ns)
            
            current_pos = 0
            text_segments = []
            
            range_starts = {} # id -> start_pos
            range_ends = {}   # id -> end_pos
            
            for elem in elements:
                tag = elem.tag
                if tag.endswith('commentRangeStart'):
                    cid = elem.get(f"{{{ns['w']}}}id")
                    if cid is not None:
                        range_starts[cid] = current_pos
                elif tag.endswith('commentRangeEnd'):
                    cid = elem.get(f"{{{ns['w']}}}id")
                    if cid is not None:
                        range_ends[cid] = current_pos
                elif tag.endswith('t'):
                    text = elem.text or ""
                    text_segments.append(text)
                    current_pos += len(text)

            full_text = "".join(text_segments)
            
            for cid, ctext in comments_map.items():
                if cid in range_starts and cid in range_ends:
                    start = range_starts[cid]
                    end = range_ends[cid]
                    
                    highlighted_text = full_text[start:end]
                    
                    # Get 100 words before
                    text_before = full_text[:start]
                    words_before = text_before.split()
                    context_before = " ".join(words_before[-100:])
                    
                    # Get 100 words after
                    text_after = full_text[end:]
                    words_after = text_after.split()
                    context_after = " ".join(words_after[:100])
                    
                    # Combine for "report_text" as requested: "get the highlighed text and add 100 words previous and 100 following workds"
                    # The user said: "report text meaning the original content that is commented." 
                    # BUT ALSO: "get the highlighed text and add 100 words previous and 100 following workds to create the text."
                    # I will construct a `full_context_text`
                    
                    report_text = f"{context_before}\n\n>> {highlighted_text} <<\n\n{context_after}"

                    comments.append({
                        "id": cid,
                        "comment": ctext,
                        "report_text": report_text, # The full context window
                        "highlighted_text": highlighted_text # The specific text commented on
                    })

    except Exception as e:
        print(f"Error extracting comments: {e}")
        return []

    return comments

def convert_markdown_to_docx(markdown_content: str) -> BytesIO:
    """
    Converts markdown content to a DOCX file in memory.
    """
    document = Document()
    
    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = markdown_content.split('\n')
    
    in_code_block = False
    code_block_content = []

    for line in lines:
        stripped_line = line.strip()

        # Handle Code Blocks
        if stripped_line.startswith('```'):
            if in_code_block:
                # End of code block
                in_code_block = False
                p = document.add_paragraph()
                p.style = 'No Spacing'
                runner = p.add_run('\n'.join(code_block_content))
                runner.font.name = 'Courier New'
                runner.font.size = Pt(9)
                code_block_content = []
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue

        # Handle Headers
        if stripped_line.startswith('#'):
            level = len(stripped_line.split()[0])
            text = stripped_line.lstrip('#').strip()
            # python-docx supports levels 1-9, but let's cap at 3 for simplicity or map
            # standard H1 -> Heading 1
            if level > 9: level = 9
            document.add_heading(text, level=level)
            continue
            
        # Handle Horizontal Rules
        if stripped_line == '---' or stripped_line == '***':
            document.add_paragraph('_' * 40) # Simple visual separator
            continue

        # Handle Lists
        if stripped_line.startswith('- ') or stripped_line.startswith('* '):
            text = stripped_line[2:].strip()
            p = document.add_paragraph(text, style='List Bullet')
            _format_runs(p)
            continue
            
        if re.match(r'^\d+\.\s', stripped_line):
            # Ordered list
            parts = stripped_line.split('.', 1)
            text = parts[1].strip()
            p = document.add_paragraph(text, style='List Number')
            _format_runs(p)
            continue

        # Normal Paragraph (skip empty lines if desired, or keep them)
        if not stripped_line:
            # document.add_paragraph("") # Adds empty line
            continue

        p = document.add_paragraph(stripped_line)
        _format_runs(p)

    # Save to memory
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def _format_runs(paragraph):
    """
    Simple helper to apply bold/italic formatting to a paragraph.
    Note: This is a basic implementation and doesn't handle nested overlapping tags perfectly.
    """
    # This is tricky in python-docx because you have to rebuild runs.
    # For a robust solution, we'd need a parser.
    # Here is a very simplified approach: check for **bold** and *italic*
    
    text = paragraph.text
    # Check if there is any formatting needed
    if "**" not in text and "*" not in text and "_" not in text and "`" not in text:
        return

    # Clear existing text and rebuild
    paragraph.clear()
    
    # Regex to split by formatting tokens: **bold**, *italic*, `code`
    # We prioritize Bold, then Italic, then Code
    
    # Tokenizer pattern: capture (**.*?**)|(\*.*?\*)|(`.*?`)
    # This splits the string, keeping delimiters
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for token in tokens:
        if not token:
            continue
            
        run = paragraph.add_run()
        
        if token.startswith('**') and token.endswith('**') and len(token) > 4:
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith('*') and token.endswith('*') and len(token) > 2:
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith('`') and token.endswith('`') and len(token) > 2:
            run.text = token[1:-1]
            run.font.name = 'Courier New'
        else:
            run.text = token
